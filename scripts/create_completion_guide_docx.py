from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "DT_Lab_Website_Form_Completion_Guide.docx"

document = Document()
section = document.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.8)
section.right_margin = Inches(0.8)

styles = document.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.5)
styles["Title"].font.name = "Aptos Display"
styles["Title"].font.size = Pt(24)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(23, 54, 93)
styles["Heading 1"].font.color.rgb = RGBColor(23, 54, 93)
styles["Heading 2"].font.color.rgb = RGBColor(68, 114, 196)

title = document.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("DT Lab Website Information Form")
subtitle = document.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A friendly guide for lab members")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(89, 89, 89)

intro = document.add_paragraph()
intro.add_run("Thank you for helping us keep the DT Lab website accurate and up to date. ").bold = True
intro.add_run(
    "The workbook already contains information currently shown or planned for the website. "
    "Please review the information relevant to you, correct anything that is inaccurate, and add any missing details. "
    "You only need to complete the sections that apply to you."
)


def heading(text, level=1):
    document.add_heading(text, level=level)


def bullets(items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


heading("Before you begin")
bullets([
    "Open DT_Lab_Website_Information_Collection_Form.xlsx and read the Start Here worksheet.",
    "Only provide information that you are comfortable publishing on the public lab website.",
    "Enter your name in Completed / confirmed by so we know whom to contact if clarification is needed.",
    "If you are unsure about a field, leave it blank and add a short explanation in Notes.",
])

heading("Status options")
table = document.add_table(rows=1, cols=2)
table.style = "Light Shading Accent 1"
table.rows[0].cells[0].text = "Status"
table.rows[0].cells[1].text = "When to use it"
statuses = [
    ("Please review", "The information still needs to be checked."),
    ("Confirmed", "The current information is accurate and may be published."),
    ("Updated", "You have corrected or replaced information in the row."),
    ("New entry", "This information is not yet on the website."),
    ("Do not publish", "The information should not appear on the public website."),
    ("Not applicable", "The field or row does not apply."),
]
for status, meaning in statuses:
    cells = table.add_row().cells
    cells[0].text = status
    cells[1].text = meaning

heading("How to complete each worksheet")

heading("1. Lab Details", 2)
document.add_paragraph(
    "This worksheet is mainly for lab leaders or administrators. Review the lab name, description, affiliation, "
    "public contact details, postal address, and online profiles. Enter the final wording in "
    "Your confirmed or updated content."
)

heading("2. Lab Members", 2)
document.add_paragraph("Find your existing row or add a new row at the bottom. Please check:")
bullets([
    "Full name, preferred title, and lab role",
    "Affiliation, research interests, and short biography",
    "Public email and professional profile links",
    "Location and profile photo file name",
])
document.add_paragraph(
    "Required fields are marked in the headings. Optional fields may be left blank. "
    "Keep research interests concise. A two- or three-sentence biography written in the third person is usually enough."
)
p = document.add_paragraph()
p.add_run("Photo guidance: ").bold = True
p.add_run(
    "Use a clear JPG or PNG portrait. Square or portrait orientation is preferred. "
    "Enter the exact file name in the workbook and submit the image with the form."
)

heading("3. Research Themes", 2)
document.add_paragraph(
    "This worksheet is mainly for lab leaders and research theme coordinators. Confirm each theme title and provide "
    "a clear 25–60 word description for a broad academic audience. Add representative keywords and indicate whether "
    "the theme should appear on the home page."
)

heading("4. Publications", 2)
document.add_paragraph("Add one publication per row. Copy details from the official publication record whenever possible:")
bullets([
    "Full title and authors in published order",
    "Conference or journal, year, and publication type",
    "DOI, arXiv, PDF, code, and project links where available",
])
document.add_paragraph(
    "A summary is optional. If included, explain the main contribution in one or two factual sentences."
)

heading("5. News", 2)
document.add_paragraph(
    "Add one item per row. Suitable items include paper acceptances, awards, grants, new members, invited talks, "
    "events, and dataset or software releases. Use YYYY-MM-DD for dates, keep headlines short, and write concise, "
    "factual descriptions. Add a supporting link or image when available."
)

heading("Writing and formatting tips")
bullets([
    "Use complete URLs beginning with https:// or http://.",
    "Enter email addresses as name@example.edu.",
    "Use YYYY-MM-DD for dates, for example 2026-07-31.",
    "Use official spellings for names, venues, organisations, and titles.",
    "Keep wording concise, factual, and suitable for a public academic website.",
    "Do not include confidential, sensitive, or unpublished information that should not be public.",
])

heading("Final checklist")
bullets([
    "Your information is accurate and approved for public use.",
    "Required fields are complete.",
    "Each reviewed row has an appropriate status.",
    "Your name appears in Completed / confirmed by.",
    "Photos and supporting files are attached and match the file names entered.",
    "Links work correctly.",
    "Any uncertainty or special request is explained in Notes.",
])

closing = document.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = closing.add_run("Thank you for contributing to the DT Lab website.")
run.bold = True
run.font.color.rgb = RGBColor(23, 54, 93)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("DT Lab Website Information Form — Member Guide").font.size = Pt(9)

document.save(OUTPUT)
print(OUTPUT)
